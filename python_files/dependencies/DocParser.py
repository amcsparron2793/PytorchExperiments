"""
*** Word_Scraper ***

uses python-docx which needs to be imported as docx
Searches for FIDs in a given word doc
if there are FIDs found, then this sends them to Main
once main completes this adds them to a csv (?) of processed FIDs with date etc

ChangeLog:

9/4/2020 AJM - Split get_text into two functions, get_filename and get_text

9/18/2020 AJM - fixed entered filename doesn't exist -
checks for existence and if False returns filename as None
an if statement from Main then reruns the function

started error handling for no FID found in Doc, still in progress
"""

# TODO: make a way to iterate through multiple documents automatically -
#  walk the input dir to get filenames - this part works
#  add the names to a list and run the whole thing for each file?

# imports
from os.path import isfile, join
from os import listdir
# installs as python-docx
import docx
from sys import exit
from time import sleep
import dependencies.CustomLog_Classes as Clog
import questionary


class DocParser:
    """ Allows Reading and Parsing docx files. """
    def __init__(self):
        self.err = Clog.Error()
        self.err.error_setup()
        self.filename, self.filepath, self.fullpath = self.GetFileNamePath()
        self.doc = self.GetDocObj()

    def GetDocObj(self):
        obj = docx.Document(self.fullpath)
        return obj

    def GetFileNamePath(self):
        while True:
            try:
                fullpath = questionary.path("Please enter path to docx file").ask()
                if isfile(fullpath) and (fullpath.endswith(".doc")
                                         or fullpath.endswith(".docx")):
                    filename = fullpath.split("/"[-1])
                    filepath = '/'.join(fullpath.split("/")[:-1])
                    print(filename, filepath)
                    return filename, filepath, fullpath
                else:
                    print("path must be to a doc or docx file that already exists")
            except questionary.ValidationError as e:
                self.err.error_handle(e)

    def GetFulltext(self):
        """ Get the full text of a docx as one large string. """
        try:
            # create a list for the fulltext (each line of the document is an entry)
            fulltext = []

            # for each line in doc, append it to the list, replacing any double tab characters with newline chars
            for para in self.doc.paragraphs:
                fulltext.append(para.text.replace('\t\t', '\n'))
            # return the fulltext as one big string joined with newlines
            return '\n'.join(fulltext)

        except Exception as e:
            self.err.error_handle(e)

    def get_text(self, string_to_find):
        """ find any given string in a docx."""
        try:
            # init counter
            i = 0

            # for each line in the doc obj increment the counter.
            # If string_to_find is in the line.
            # Print it and print the line number it was found on.
            for para in self.doc.paragraphs:
                i += 1
                if string_to_find in para.text:
                    print(para.text)
                    print('line number is {paragraph_num}'.format(paragraph_num=str(i)))

        except Exception as e:
            self.err.error_handle_no_exit_quiet(e)

    # deprecated
    def get_filenames(self, from_path=None):
        docx_found = []

        if from_path is None:
            misc_proj_files = join(Clog.project_root, 'Misc_Project_Files')
            from_path = misc_proj_files

        for doc in listdir(from_path):
            if doc.endswith('.docx') or doc.endswith('.doc'):
                docx_found.append(doc)
            else:
                pass
        try:
            if len(docx_found) > 0:
                print('Detected docx\'s: ' + str(docx_found))
                while True:
                    filename = input('Enter filename or '
                                     'press Q to quit: ').lower()

                    # if q is entered then quit
                    if filename == 'q':
                        print('quitting....')
                        # this is imported using a from statement so 'time.' is not needed
                        sleep(1)
                        # this was imported using a from statement so 'sys.' is not needed
                        exit()

                    elif isfile(join(from_path, filename)):
                        full_path = join(from_path, filename)
                        return full_path, filename

                    elif not isfile(join(from_path, filename)):
                        print("{fname} not found in {search_loc}. Please try again".format(fname=filename,
                                                                                           search_loc=from_path))

                    else:
                        print('Invalid entry, please try again.')
            elif not len(docx_found) > 0:
                raise FileNotFoundError('No Files found in {fpath}'.format(fpath=from_path))
        except FileNotFoundError as e:
            self.err.error_handle(e)
        except Exception as e:
            self.err.error_handle(e)
